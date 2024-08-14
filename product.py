import streamlit as st
from groq import Groq
import os
import base64
# from google.oauth2.credentials import Credentials
# from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from google.auth.transport.requests import Request
from datetime import datetime, timedelta
from email.mime.text import MIMEText
import pdfplumber
from io import BytesIO
from docx import Document
import pandas as pd
import json
import imaplib
import email
from email.message import EmailMessage
import smtplib
import gspread
from google.oauth2.service_account import Credentials
from email.utils import parsedate_to_datetime



# Set up the necessary scopes and credentials file
SCOPES = ['https://www.googleapis.com/auth/gmail.readonly', 'https://www.googleapis.com/auth/gmail.send','https://www.googleapis.com/auth/gmail.modify']
GOOGLE_CREDENTIALS_JSON = st.secrets['GOOGLE_CREDENTIALS_JSON']

def get_gmail_service():
    if 'creds' not in st.session_state:
        st.session_state.creds = None

    creds = st.session_state.creds

    if creds:
        creds = Credentials.from_authorized_user_info(json.loads(creds), SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            credentials_info = json.loads(GOOGLE_CREDENTIALS_JSON)
            flow = InstalledAppFlow.from_client_config(credentials_info, SCOPES, redirect_uri='https://sopappuct-7otfy47eudpdpdw4cpyrhe.streamlit.app/')
            # st.write("1")
            auth_url, _ = flow.authorization_url(prompt='consent')
            st.session_state.auth_url = auth_url
            if st.session_state.Authenticate_check:
                st.sidebar.write(st.session_state.auth_url)
            # st.write("1")
            auth_code = st.experimental_get_query_params().get('code')
            if auth_code:
                auth_code = auth_code[0]
                flow.fetch_token(code=auth_code)
                creds = flow.credentials
                st.session_state.creds = creds.to_json()

        if creds:
            st.session_state.creds = creds.to_json()
    # st.write("1")
    service = build('gmail', 'v1', credentials=creds)
    return service

def get_message_details(message):
    """Extracts the subject and content of an email message."""
    headers = message['payload']['headers']
    subject = next(header['value'] for header in headers if header['name'] == 'Subject')
    
    if 'data' in message['payload']['body']:
        body_data = message['payload']['body']['data']
        body_decoded = base64.urlsafe_b64decode(body_data).decode('utf-8')
        return subject, body_decoded
    else:
        parts = message['payload'].get('parts', [])
        for part in parts:
            if part['mimeType'] == 'text/plain':
                body_data = part['body'].get('data', '')
                body_decoded = base64.urlsafe_b64decode(body_data).decode('utf-8')
                return subject, body_decoded
            elif part['mimeType'] == 'text/html':
                body_data = part['body'].get('data', '')
                body_decoded = base64.urlsafe_b64decode(body_data).decode('utf-8')
                return subject, body_decoded
    return subject, "No content available"

# def remove_prefix(subject):
#     prefixes = ["Re:", "Fwd:", "Fw:", "RE:", "FWD:", "FW:"]
#     for prefix in prefixes:
#         if subject.startswith(prefix):
#             return subject[len(prefix):].strip()
#     return subject.strip()

# def count_email():
#     imap_server = 'imap.gmail.com'
#     email_user = st.session_state.user_gmail
#     email_pass = st.session_state.password
    
#     mail = imaplib.IMAP4_SSL(imap_server)
#     mail.login(email_user, email_pass)
#     mail.select('inbox')  # Select the mailbox you want to use
    
#     # Search for emails from a specific sender
#     status, messages = mail.search(None, f'FROM "{st.session_state.gmail_sender}"')
    
#     # Get the list of email IDs
#     email_ids = messages[0].split()
    
#     # Use a set to store unique subjects
#     unique_subjects = set()
    
#     # Function to remove common prefixes
    
#     for email_id in email_ids:
#         status, data = mail.fetch(email_id, '(BODY.PEEK[HEADER])')
#         msg = email.message_from_bytes(data[0][1])
#         subject = msg.get('Subject')
#         if subject:
#             # Normalize the subject by removing common prefixes
#             normalized_subject = remove_prefix(subject)
#             unique_subjects.add(normalized_subject)
    
#     # Count the number of unique subjects
#     return len(unique_subjects)

# # Function to fetch the latest email from a sender
# def fetch_latest_email():
#     if st.button("Fetch Gmail"):
#         if st.session_state.password and st.session_state.gmail_sender and st.session_state.user_gmail:
#             imap_server = "imap.gmail.com"
            
#             # Connect to the IMAP server
#             mail = imaplib.IMAP4_SSL(imap_server)
#             try:
#                 mail.login(st.session_state.user_gmail, st.session_state.password)
#             except Exception as e:
#                 st.error("invalid username or password")
#             mail.select("inbox")
            
#             # Search for emails from the specific sender
#             status, data = mail.search(None, f'FROM "{st.session_state.gmail_sender}"')
            
#             email_ids = data[0].split()
#             if not email_ids:
#                 mail.logout()
#                 return None, None, None
            
#             latest_email_id = email_ids[-1]
#             status, msg_data = mail.fetch(latest_email_id, "(RFC822)")
#             raw_email = msg_data[0][1]
#             original_email = email.message_from_bytes(raw_email)
            
#             mail.logout()
            
#             # return original_email, latest_email_id, raw_email
        
#         # # Function to display email content
#         # def display_email_content(original_email):
#             st.session_state.fetched_subject = original_email["Subject"]
#             st.session_state.fetched_sender_gmail = original_email["From"]
#             st.session_state.msg_id = original_email["Message-ID"]
#             st.session_state.email_count_total = count_email()
#             st.write(f"Total Request is :{st.session_state.email_count_total}")
#             st.write(st.session_state.msg_id)
            
#             st.write(f"**Subject:** {st.session_state.fetched_subject}")
#             st.write(f"**From:** {st.session_state.fetched_sender_gmail}")
            
#             # Extract and display the body of the email
#             email_body = ""
#             if original_email.is_multipart():
#                 for part in original_email.walk():
#                     if part.get_content_type() == "text/plain":
#                         st.session_state.fetched_content = part.get_payload(decode=True).decode()
#                         break
#             else:
#                 st.session_state.fetched_content = original_email.get_payload(decode=True).decode()
            
#             st.write(f"**Body:**\n{st.session_state.fetched_content}")
#         else:
#             st.error("Please fill all the field")
    
#     # return original_subject, original_from, original_message_id, email_body







# def remove_prefix(subject):
#     prefixes = ["Re:", "Fwd:", "Fw:", "RE:", "FWD:", "FW:"]
#     for prefix in prefixes:
#         if subject.startswith(prefix):
#             return subject[len(prefix):].strip()
#     return subject.strip()

# Function to count unique subjects from a specific sender
def count_email():
    imap_server = 'imap.gmail.com'
    email_user = st.session_state.user_gmail
    email_pass = st.session_state.password
    
    mail = imaplib.IMAP4_SSL(imap_server)
    mail.login(email_user, email_pass)
    mail.select('inbox')  # Select the mailbox you want to use
    
    # Search for emails from a specific sender
    status, messages = mail.search(None, f'FROM "{st.session_state.gmail_sender}"')
    
    # Get the list of email IDs
    email_ids = messages[0].split()
    
    # Use a set to store unique subjects
    unique_subjects = set()
    
    for email_id in email_ids:
        status, data = mail.fetch(email_id, '(BODY.PEEK[HEADER])')
        if status != 'OK':
            continue
        msg = email.message_from_bytes(data[0][1])
        subject = msg.get('Subject')
        if subject:
            # Normalize the subject by removing common prefixes
            normalized_subject = remove_prefix(subject)
            unique_subjects.add(normalized_subject)
    
    # Close the connection
    mail.logout()

    # Count the number of unique subjects
    return len(unique_subjects)

# # Function to fetch the latest email from a sender
# def fetch_latest_email():
#     if st.button("Fetch Gmail"):
#         if st.session_state.password and st.session_state.gmail_sender and st.session_state.user_gmail:
#             imap_server = "imap.gmail.com"
            
#             # Connect to the IMAP server
#             mail = imaplib.IMAP4_SSL(imap_server)
#             try:
#                 mail.login(st.session_state.user_gmail, st.session_state.password)
#             except Exception as e:
#                 st.error("Invalid username or password")
#                 return
            
#             mail.select("inbox")
            
#             # Search for emails from the specific sender
#             status, data = mail.search(None, f'FROM "{st.session_state.gmail_sender}"')
            
#             email_ids = data[0].split()
#             if not email_ids:
#                 st.error("No emails found from the specified sender.")
#                 mail.logout()
#                 return
            
#             latest_email_id = email_ids[-1]
#             status, msg_data = mail.fetch(latest_email_id, "(RFC822)")
#             if status != 'OK':
#                 st.error("Failed to fetch the latest email.")
#                 mail.logout()
#                 return
            
#             raw_email = msg_data[0][1]
#             original_email = email.message_from_bytes(raw_email)
            
#             # Close the connection
#             mail.logout()
            
#             # Display the email content
#             st.session_state.fetched_subject = original_email["Subject"]
#             st.session_state.fetched_sender_gmail = original_email["From"]
#             st.session_state.msg_id = original_email["Message-ID"]
            # st.session_state.email_count_total = count_email()
            # st.write(f"Total Request is: {st.session_state.email_count_total}")
#             # st.write(f"Message-ID: {st.session_state.msg_id}")
            
#             st.write(f"**Subject:** {st.session_state.fetched_subject}")
#             st.write(f"**From:** {st.session_state.fetched_sender_gmail}")
            
#             # Extract and display the body of the email
#             if original_email.is_multipart():
#                 for part in original_email.walk():
#                     if part.get_content_type() == "text/plain":
#                         st.session_state.fetched_content = part.get_payload(decode=True).decode()
#                         break
#             else:
#                 st.session_state.fetched_content = original_email.get_payload(decode=True).decode()
            
#             st.write(f"**Body:**\n{st.session_state.fetched_content}")
#         else:
#             st.error("Please fill in all the fields")




# def remove_prefix(subject):
#     prefixes = ["Re:", "Fwd:", "Fw:", "RE:", "FWD:", "FW:"]
#     for prefix in prefixes:
#         if subject.startswith(prefix):
#             return subject[len(prefix):].strip()
#     return subject.strip()

# def get_email_body(original_email):
#     if original_email.is_multipart():
#         for part in original_email.walk():
#             if part.get_content_type() == "text/plain":
#                 return part.get_payload(decode=True).decode()
#     else:
#         return original_email.get_payload(decode=True).decode()
#     return ""

# def fetch_latest_email():
#     if st.button("Fetch Gmail"):
#         if st.session_state.password and st.session_state.gmail_sender and st.session_state.user_gmail:
#             st.session_state.email_count_total = count_email()
            
#             imap_server = "imap.gmail.com"
            
#             # Connect to the IMAP server
#             mail = imaplib.IMAP4_SSL(imap_server)
#             try:
#                 mail.login(st.session_state.user_gmail, st.session_state.password)
#             except Exception as e:
#                 st.error("Invalid username or password")
#                 return
            
#             mail.select("inbox")
            
#             # Search for emails from the specific sender
#             status_from, data_from = mail.search(None, f'FROM "{st.session_state.gmail_sender}"')
#             status_to, data_to = mail.search(None, f'TO "{st.session_state.gmail_sender}"')
            
#             # Combine results and remove duplicates
#             email_ids = list(set(data_from[0].split() + data_to[0].split()))
#             if not email_ids:
#                 mail.logout()
#                 st.error("No emails found from the specified sender")
#                 return
            
#             st.session_state.emails = []
            
#             for email_id in email_ids:
#                 status, msg_data = mail.fetch(email_id, "(RFC822)")
#                 raw_email = msg_data[0][1]
#                 original_email = email.message_from_bytes(raw_email)
#                 subject = remove_prefix(original_email['Subject'])
#                 st.session_state.emails.append({
#                     'subject': subject,
#                     'from': original_email['From'],
#                     'message_id': original_email['Message-ID'],
#                     'in_reply_to': original_email['In-Reply-To'],
#                     'references': original_email['References'],
#                     'body': get_email_body(original_email),
#                     'date': original_email['Date']
#                 })
            
#             mail.logout()
#         else:
#             st.error("Please fill all the fields")

def remove_prefix(subject):
    prefixes = ["Re:", "Fwd:", "Fw:", "RE:", "FWD:", "FW:"]
    subject = subject.strip()  # Remove leading and trailing spaces
    while True:
        for prefix in prefixes:
            if subject.startswith(prefix):
                subject = subject[len(prefix):].strip()
                break
        else:
            break  # Exit the loop if no prefix is found
    return subject


def get_email_body(original_email):
    if original_email.is_multipart():
        for part in original_email.walk():
            if part.get_content_type() == "text/plain":
                return part.get_payload(decode=True).decode()
    else:
        return original_email.get_payload(decode=True).decode()
    return ""

# def fetch_latest_email():
#     if st.button("Fetch Gmail"):
#         if st.session_state.password and st.session_state.gmail_sender and st.session_state.user_gmail:
#             st.session_state.email_count_total = count_email()
            
#             imap_server = "imap.gmail.com"
            
#             # Connect to the IMAP server
#             mail = imaplib.IMAP4_SSL(imap_server)
#             try:
#                 mail.login(st.session_state.user_gmail, st.session_state.password)
#             except Exception as e:
#                 st.error("Invalid username or password")
#                 return
            
#             mail.select("inbox")
            
#             # Search for emails from the specific sender
#             status_from, data_from = mail.search(None, f'FROM "{st.session_state.gmail_sender}"')
#             status_to, data_to = mail.search(None, f'TO "{st.session_state.gmail_sender}"')
            
#             # Combine results and remove duplicates
#             email_ids = list(set(data_from[0].split() + data_to[0].split()))
#             if not email_ids:
#                 mail.logout()
#                 st.error("No emails found from the specified sender")
#                 return
            
#             st.session_state.emails = []
            
#             for email_id in email_ids:
#                 status, msg_data = mail.fetch(email_id, "(RFC822)")
#                 raw_email = msg_data[0][1]
#                 original_email = email.message_from_bytes(raw_email)
#                 subject = remove_prefix(original_email['Subject'])
#                 date_str = original_email['Date']
#                 date = parsedate_to_datetime(date_str) if date_str else None
#                 st.session_state.emails.append({
#                     'subject': subject,
#                     'from': original_email['From'],
#                     'message_id': original_email['Message-ID'],
#                     'in_reply_to': original_email['In-Reply-To'],
#                     'references': original_email['References'],
#                     'body': get_email_body(original_email),
#                     'date': date
#                 })
            
#             mail.logout()
#         else:
#             st.error("Please fill all the fields")

def fetch_latest_email():
    if st.button("Fetch Gmail"):
        if st.session_state.password and st.session_state.gmail_sender and st.session_state.user_gmail:
            st.session_state.email_count_total = count_email()
            
            imap_server = "imap.gmail.com"
        
            # Connect to the IMAP server
            mail = imaplib.IMAP4_SSL(imap_server)
            try:
                mail.login(st.session_state.user_gmail, st.session_state.password)
            except Exception as e:
                st.error("Invalid username or password")
                return
            
            # Select the "Inbox" label to search for received emails
            mail.select("inbox")
            
            # Search for emails from and to the specific recipient
            status_from, data_from = mail.search(None, f'FROM "{st.session_state.gmail_sender}"')
            status_to, data_to = mail.search(None, f'TO "{st.session_state.gmail_sender}"')
            
            # Combine results and remove duplicates
            received_email_ids = data_from[0].split()
            sent_email_ids = data_to[0].split()
            email_ids = list(set(received_email_ids + sent_email_ids))
            
            if not email_ids:
                mail.logout()
                st.error("No emails found from or to the specified recipient")
                return
            
            st.session_state.emails = []
            for email_id in email_ids:
                status, msg_data = mail.fetch(email_id, "(RFC822)")
                raw_email = msg_data[0][1]
                original_email = email.message_from_bytes(raw_email)
                subject = remove_prefix(original_email['Subject'])
                date_str = original_email['Date']
                date = parsedate_to_datetime(date_str) if date_str else None
                st.session_state.emails.append({
                    'subject': subject,
                    'from': original_email['From'],
                    'to': original_email.get('To', 'N/A'),
                    'message_id': original_email['Message-ID'],
                    'in_reply_to': original_email['In-Reply-To'],
                    'references': original_email['References'],
                    'body': get_email_body(original_email),
                    'date': date
                })
            
            mail.logout()
        else:
            st.error("Please fill all the fields")

def fetch_received_emails():
    if st.session_state.password and st.session_state.gmail_sender and st.session_state.user_gmail:
        # st.session_state.email_count_total = count_email()
        imap_server = "imap.gmail.com"
        
        # Connect to the IMAP server
        mail = imaplib.IMAP4_SSL(imap_server)
        try:
            mail.login(st.session_state.user_gmail, st.session_state.password)
        except Exception as e:
            st.error("Invalid username or password")
            return
        
        # Select the "Inbox" label to search for received emails
        mail.select("inbox")
        
        # Search for emails from the specific sender
        status_from, data_from = mail.search(None, f'FROM "{st.session_state.gmail_sender}"')
        
        email_ids = data_from[0].split()
        if not email_ids:
            mail.logout()
            st.error("No received emails found from the specified sender")
            return
        
        st.session_state.received_emails = []
        for email_id in email_ids:
            status, msg_data = mail.fetch(email_id, "(RFC822)")
            raw_email = msg_data[0][1]
            original_email = email.message_from_bytes(raw_email)
            subject = remove_prefix(original_email['Subject'])
            date_str = original_email['Date']
            date = None
            if date_str:
                try:
                    # Parse the date and convert it to UTC if it's not already in UTC
                    date = parsedate_to_datetime(date_str)
                    if date.tzinfo is None:
                        # Assume the time is in UTC if no time zone info is provided
                        date = date.replace(tzinfo=timezone.utc)
                    else:
                        # Convert to UTC for consistency
                        date = date.astimezone(timezone.utc)
                except Exception as e:
                    st.error(f"Error parsing date: {e}")
            # date = parsedate_to_datetime(date_str) if date_str else None
            st.session_state.received_emails.append({
                'subject': subject,
                'from': original_email['From'],
                'to': original_email.get('To', 'N/A'),
                'message_id': original_email['Message-ID'],
                'in_reply_to': original_email['In-Reply-To'],
                'references': original_email['References'],
                'body': get_email_body(original_email),
                'date': date
            })
        
        mail.logout()
    else:
        st.error("Please fill all the fields")

def fetch_sent_emails():
    if st.session_state.password and st.session_state.gmail_sender and st.session_state.user_gmail:
        imap_server = "imap.gmail.com"
        
        # Connect to the IMAP server
        mail = imaplib.IMAP4_SSL(imap_server)
        try:
            mail.login(st.session_state.user_gmail, st.session_state.password)
        except Exception as e:
            st.error("Invalid username or password")
            return
        
        # Select the "Sent" label to search for sent emails
        mail.select('"[Gmail]/Sent Mail"')
        
        # Search for emails sent to the specific recipient
        status, data = mail.search(None, f'TO "{st.session_state.gmail_sender}"')
        
        email_ids = data[0].split()
        if not email_ids:
            mail.logout()
            st.error("No sent emails found to the specified recipient")
            return
        
        st.session_state.sent_emails = []
        for email_id in email_ids:
            status, msg_data = mail.fetch(email_id, "(RFC822)")
            raw_email = msg_data[0][1]
            original_email = email.message_from_bytes(raw_email)
            subject = remove_prefix(original_email['Subject'])
            date_str = original_email['Date']
            date = None
            if date_str:
                try:
                    # Parse the date and convert it to UTC if it's not already in UTC
                    date = parsedate_to_datetime(date_str)
                    if date.tzinfo is None:
                        # Assume the time is in UTC if no time zone info is provided
                        date = date.replace(tzinfo=timezone.utc)
                    else:
                        # Convert to UTC for consistency
                        date = date.astimezone(timezone.utc)
                except Exception as e:
                    st.error(f"Error parsing date: {e}")
            # date = parsedate_to_datetime(date_str) if date_str else None
            st.session_state.sent_emails.append({
                'subject': subject,
                'from': original_email['From'],
                'to': original_email.get('To', 'N/A'),
                'message_id': original_email['Message-ID'],
                'in_reply_to': original_email['In-Reply-To'],
                'references': original_email['References'],
                'body': get_email_body(original_email),
                'date': date
            })
        
        mail.logout()
    else:
        st.error("Please fill all the fields")

# def combine_emails():
#     # Combine received and sent emails
#     if 'received_emails' in st.session_state and 'sent_emails' in st.session_state:
#         st.session_state.emails = st.session_state.received_emails + st.session_state.sent_emails
#     elif 'received_emails' in st.session_state:
#         st.session_state.emails = st.session_state.received_emails
#     elif 'sent_emails' in st.session_state:
#         st.session_state.emails = st.session_state.sent_emails
#     else:
#         st.session_state.emails = []
# def combine_emails():
#     if 'received_emails' in st.session_state and 'sent_emails' in st.session_state:
#         received_subjects = {email['subject'] for email in st.session_state.received_emails}
#         sent_emails_to_include = [email for email in st.session_state.sent_emails if email['subject'] in received_subjects]
#         st.session_state.emails = st.session_state.received_emails + sent_emails_to_include
#     elif 'received_emails' in st.session_state:
#         st.session_state.emails = st.session_state.received_emails
#     elif 'sent_emails' in st.session_state:
#         st.session_state.emails = []  # No matching subjects to include sent emails alone
#     else:
#         st.session_state.emails = []
def combine_emails():
    if 'received_emails' not in st.session_state:
        st.session_state.received_emails = []
    if 'sent_emails' not in st.session_state:
        st.session_state.sent_emails = []

    st.session_state.emails = []

    # Create a dictionary of subjects from received emails
    received_subjects = {email['subject']: email for email in st.session_state.received_emails}

    # Combine sent emails with matching subjects in received emails
    for sent_email in st.session_state.sent_emails:
        if sent_email['subject'] in received_subjects:
            st.session_state.emails.append(sent_email)

    # Add the received emails that were not replied to
    st.session_state.emails.extend(st.session_state.received_emails)

    # Sort the emails by date in ascending order
    st.session_state.emails.sort(key=lambda e: e['date'] if e['date'] else datetime.min)


def filter_emails(filter_option):
    if filter_option == "Pending Request":
        return [email for email in st.session_state.received_emails if email['message_id'] not in {e['in_reply_to'] for e in st.session_state.sent_emails}]
    elif filter_option == "Replied Request":
        sent_message_ids = {email['message_id'] for email in st.session_state.sent_emails}
        return [email for email in st.session_state.received_emails if email['in_reply_to'] in sent_message_ids]
    else:
        return st.session_state.emails

def display_thread(selected_email_index):
    email_threads = {}
    for email_data in st.session_state.emails:
        email_threads[email_data['message_id']] = email_data

    current_email = st.session_state.emails[selected_email_index]
    thread = []

    # Traverse replies in chronological order
    while current_email:
        thread.append(current_email)
        in_reply_to_id = current_email['in_reply_to']
        if in_reply_to_id:
            current_email = email_threads.get(in_reply_to_id, None)
        else:
            current_email = None
    
    for email_data in reversed(thread):
        st.write(f"**Date:** {email_data['date']}")
        st.write(f"**Subject:** {email_data['subject']}")
        st.write(f"**From:** {email_data['from']}")
        st.write(f"**To:** {email_data['to']}")
        st.write(f"**Body:**\n{email_data['body']}")
        st.write("---")
# Function to send a reply email
def send_reply_email():
    st.title('Send an Email via Gmail')
    st.write('Enter the details below to send an email.')

    sender_email = st.text_input('Employee Email Address',st.session_state.user_gmail)
    recipient_email = st.text_input('Client Email Address',st.session_state.gmail_sender)
    Cc = st.text_input('Cc Email Address')
    password = st.text_input("Password", type="password")
    st.session_state.user_subject = st.text_input('Subject',st.session_state.user_subject)
    subject = st.session_state.user_subject
    st.session_state.content = st.text_area('Message',st.session_state.content,height=500)
    message_text = st.session_state.content
    if st.button('Send Email',key = 'process_end'):
        
        msg = EmailMessage()
        msg["Subject"] = subject
        msg["From"] = sender_email
        msg["To"] = recipient_email
        msg['Cc'] = Cc
        msg["In-Reply-To"] = st.session_state.msg_id
        msg["References"] = st.session_state.msg_id
        
        msg.set_content(st.session_state.content)
        
        smtp_server = smtplib.SMTP("smtp.gmail.com", 587)
        smtp_server.starttls()
        smtp_server.login(st.session_state.user_gmail, password)
        smtp_server.send_message(msg)
        smtp_server.quit()
        
        st.success(f"Email sent to {st.session_state.fetched_sender_gmail}")

def fetch_gmail(sender_email):

    date = st.date_input('Date', key='date')
    message_content = ''
    subject = ''

    if st.button('Fetch Gmail Messages', key='fetch_gmail'):
        if sender_email and date:
            service = get_gmail_service()
            date_str = date.strftime('%Y/%m/%d')
            next_date_str = (date + timedelta(days=1)).strftime('%Y/%m/%d')
            query = f'from:{sender_email} after:{date_str} before:{next_date_str}'
            results = service.users().messages().list(userId='me', q=query).execute()
            messages = results.get('messages', [])

            if not messages:
                st.write('No messages found from this sender on the specified date.')
                st.session_state.gmail_fetched = False
            else:
                st.write(f'Found {len(messages)} messages from {sender_email} on {date}:')

                for msg in messages[::-1]:
                    msg_id = msg['id']
                    message = service.users().messages().get(userId='me', id=msg_id).execute()
                    snippet = message['snippet']
                    subject, message_content = get_message_details(message)
                    
                    st.write('---')
                    st.write(f'Subject: {subject}')
                    st.write(f'Message ID: {msg_id}')
                    st.write(f'Snippet: {snippet}')
                    st.write('Content:')
                    st.markdown(message_content, unsafe_allow_html=True)
                    st.write('---')
                st.session_state.msg_id = msg_id
                st.session_state.gmail_fetched = True
                st.session_state.gmail_content = f"""Subject:{subject}\nContent:{message_content}"""
                
        else:
            st.write('Please enter a sender email address and a date.')
            st.session_state.gmail_fetched = False

def create_message(sender, to, subject, message_text):
    """Create a message for an email."""
    message = MIMEText(message_text)
    message['to'] = to
    message['from'] = sender
    message['subject'] = subject
    message['In-Reply-To'] = st.session_state.msg_id
    message['References'] = st.session_state.msg_id
    raw_message = base64.urlsafe_b64encode(message.as_bytes()).decode()
    return {'raw': raw_message}

def send_message(service, user_id, message):
    """Send an email message."""
    try:
        sent_message = service.users().messages().send(userId=user_id, body=message).execute()
        st.write(f"Message sent successfully: {sent_message['id']}")
        st.session_state.gmail_send = True
    except Exception as error:
        st.write(f"An error occurred: {error}")

def create_message_for_reply(sender, to, subject, message_text, thread_id, message_id):
    message = MIMEText(message_text)
    message['to'] = to
    message['from'] = sender
    message['subject'] = subject
    message['In-Reply-To'] = message_id
    message['References'] = message_id
    raw = base64.urlsafe_b64encode(message.as_bytes()).decode('utf-8')
    return {'raw': raw, 'threadId': thread_id}

def reply_to_message(service, message_id, reply_text):
    original_message = service.users().messages().get(userId='me', id=message_id, format='full').execute()
    thread_id = original_message['threadId']
    
    headers = original_message['payload']['headers']
    for header in headers:
        if header['name'] == 'From':
            from_email = header['value']
        if header['name'] == 'Subject':
            subject = header['value']
    
    # Get the user's email address
    profile = service.users().getProfile(userId='me').execute()
    user_email = profile['emailAddress']
    
    reply_message = create_message_for_reply(user_email, from_email, subject, reply_text, thread_id, message_id)
    service.users().messages().send(userId='me', body=reply_message).execute()

def gmailsender():
    st.title('Send an Email via Gmail')
    st.write('Enter the details below to send an email.')

    sender_email = st.text_input('Sender Email Address',st.session_state.user_gmail)
    recipient_email = st.text_input('Recipient Email Address',st.session_state.gmail_sender)
    subject = st.text_input('Subject',st.session_state.subject)
    message_text = st.text_area('Message',st.session_state.content,height=500)

    if st.button('Send Email',key = 'process_end'):
        if sender_email and recipient_email and subject and message_text:
            service = get_gmail_service()
            message = create_message(sender_email, recipient_email, subject, message_text)
            send_message(service, 'me', message)
            # reply_to_message(service, st.session_state.msg_id, st.session_state.gmail_content)
            # st.success("Reply sent!")
        else:
            st.write('Please fill out all fields.')

def read_pdf(file):
    content = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            tables = page.extract_tables()
            if text:
                content.append(text)
            if tables:
                for table in tables:
                    table_text = "\n".join(["\t".join(map(str, row)) for row in table])
                    content.append(table_text)
    return "\n\n".join(content)

def read_docx(file):
    content = []
    document = Document(file)
    
    # Read paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():  # Only add non-empty paragraphs
            content.append(paragraph.text)
    
    # Read tables
    for table in document.tables:
        for row in table.rows:
            row_text = "\t".join([cell.text.strip() for cell in row.cells])
            if row_text.strip():  # Only add non-empty rows
                content.append(row_text)
    
    # Join content with two newlines, but ensure no multiple consecutive newlines
    return "\n\n".join([line for line in content if line.strip()])

def parse_feedback(text):
    feedback_part = text.split("Evaluation Based on SOP Criteria:")[0]
    evaluation_part = text.split("Evaluation Based on SOP Criteria:")[1]
    
    criteria = []
    marks = []
    reasons = []
    
    for line in evaluation_part.split('\n'):
        if line.strip():
            crit, rest = line.split(":", 1)
            mark, reason = rest.split("(", 1)
            criteria.append(crit.strip())
            marks.append(mark.strip())
            reasons.append(reason.rstrip(")").strip())
    
    return feedback_part, criteria, marks, reasons

# Function to process the input text
def process_feedback(text):
    feedback, sop_evaluation = text.split("Evaluation Based on SOP:")
    return feedback.strip(), sop_evaluation.strip()

def process_criteria(text):
    try:
        feedback, criteria = text.split("Criteria Instruction in SOP:")
    except Exception as e:
        feedback, criteria = text.split("SOP Criteria Instruction:")
    return feedback.strip(), criteria.strip()
    
# Function to convert SOP evaluation to DataFrame
def parse_sop_evaluation(sop_text):
    lines = sop_text.split('\n')
    criteria = []
    marks = []
    reasons = []

    for line in lines:
        if '|' in line and 'Criteria' not in line:
            parts = line.split('|')
            criteria.append(parts[1].strip())
            marks.append(parts[2].strip())
            reasons.append(parts[3].strip())

    data = {
        'Criteria': criteria,
        'Mark (out of 10)': marks,
        'Reason': reasons
    }

    return pd.DataFrame(data)

# Define navigation function
def navigate_to_step(step):
    st.session_state.step = step
    st.rerun()
    


def load_default_sop_file():
    # Read the content of the default file (e.g., default.txt)
    with open("default_sop_content.txt", "r") as file:
        default_content = file.read()
    return default_content

def evaluator(client):
    # st.sidebar.write("If you use gmail to fetch or send gmail please authenticate then move forward if alreadydid it ignore it")
    # if st.sidebar.button("Authenticate"):
    #     st.session_state.Authenticate_check = True
    #     st.sidebar.write("Please go to this URL to authorize the application:")
    #     get_gmail_service()
    #     st.sidebar.write(st.session_state.auth_url)

    st.sidebar.title("Navigation")
    if st.sidebar.button("Step 1: Upload SOP"):
        navigate_to_step(1)
    if st.session_state.sop_uploaded:
        if st.sidebar.button("Step 2: Client Request"):
            navigate_to_step(2)
    if st.session_state.gmail_fetched:
        if st.sidebar.button("Step 3: Evaluation and feedback"):
            navigate_to_step(3)
    if st.session_state.feedback:
        if st.sidebar.button("Step 4:Sending gmail"):
            navigate_to_step(4)
        

    if st.session_state.step == 1:
        st.title("Step 1: Upload SOP File ")
        with st.expander("Upload File", expanded=False):
            uploaded_file = st.file_uploader("Choose a file", type=["txt", "pdf", "docx"])
            if st.button("Use default SOP"):
                st.session_state.sop_content = load_default_sop_file()
        
        if uploaded_file is not None:
            if uploaded_file.name.endswith(".txt"):
                st.session_state.sop_content = uploaded_file.read().decode("utf-8")
            elif uploaded_file.name.endswith(".pdf"):
                st.session_state.sop_content = read_pdf(uploaded_file)
            elif uploaded_file.name.endswith(".docx"):
                st.session_state.sop_content = read_docx(uploaded_file)
        # else:
        #     if st.session_state.default_sop_content:
        #         st.session_state.sop_content = st.session_state.default_sop_content
        # st.session_state.sop_content = sop_content
        st.session_state.sop_content = st.text_area("Edit Your SOP Content", st.session_state.sop_content, height=300)
        col1, col3 = st.columns([1,1])
        with col3:
            if st.button("Next"):
                if st.session_state.sop_content:
                    st.session_state.sop_uploaded = True
                    navigate_to_step(2)
                else:
                    st.error("Upload SOP File or Insert SOP content.")
    
    # Step 2: Client Request
    elif st.session_state.step == 2:
        # if st.sidebar.button("Step 1: Upload SOP"):
        #     navigate_to_step(1)
        st.title("Step 2: Client Request")
        option = st.selectbox("Choose the way you want to get client request", ("Insert Client Request", "Fetch Client Request From Gmail"), index=0, placeholder='Choose an option')
        
        if option == 'Fetch Client Request From Gmail':
            st.write('Enter your gmail and the sender email address and the date to fetch your Gmail messages from that sender.')
            st.session_state.user_gmail = st.text_input('Employee Email Address',st.session_state.user_gmail, key='usergmail')
            st.session_state.gmail_sender = st.text_input('Client Email Address',st.session_state.gmail_sender, key='sender_email')
            st.session_state.password = st.text_input("Password", type="password")
            # try:
            # fetch_latest_email()
            # if st.button("Fetch Emails"):
            #     fetch_received_emails()
            # # if st.button("Fetch Sent Emails"):
            #     fetch_sent_emails()
            
            # combine_emails()
            # if st.session_state.email_count_total:
            #     st.write(f"Total Request is: {st.session_state.email_count_total}")
            
            # filter_option = st.selectbox("Filter Emails", ["All Request", "Pending Request", "Replied Request"], index=0)
    
            # filtered_emails = filter_emails(filter_option)
    
            # if filtered_emails:
            #     # Extract unique subjects
            #     unique_subjects = {}
            #     for email_data in filtered_emails:
            #         subject = email_data['subject']
            #         if subject not in unique_subjects:
            #             unique_subjects[subject] = []
            #         unique_subjects[subject].append(email_data)
    
            #     # Sort subjects by the most recent email date
            #     sorted_subjects = sorted(unique_subjects.keys(), key=lambda s: max(email['date'] for email in unique_subjects[s] if email['date']), reverse=True)
                
            #     # Display unique subjects
            #     selected_subject = st.selectbox("Select a unique request", sorted_subjects)
                
            #     # Display responses related to the selected subject
            #     related_emails = unique_subjects[selected_subject]
                
            #     # Sort related emails by date in ascending order
            #     related_emails_sorted = sorted(related_emails, key=lambda e: e['date'] if e['date'] else datetime.min)
                
            #     email_options = [f"{email['subject']} - {email['date'].strftime('%Y-%m-%d %H:%M:%S') if email['date'] else 'No Date'}" for email in related_emails_sorted]
            #     selected_email_index = st.selectbox("Select a response", email_options)
            #     selected_email = related_emails_sorted[email_options.index(selected_email_index)]
    
            #     st.session_state.fetched_subject = st.text_area("Client Subject:", selected_email['subject'], height=50)
            #     st.session_state.fetched_content = st.text_area("Client Content:", selected_email['body'], height=500)
            if st.button("Fetch Emails"):
                fetch_received_emails()
                fetch_sent_emails()
            # if st.session_state.email_count_total:
            #     st.write(f"Total Request is: {st.session_state.email_count_total}")
            st.write(st.session_state.received_emails)
            st.write(st.session_state.sent_emails)
            combine_emails()
            if 'emails' in st.session_state and st.session_state.emails:
            # Extract unique subjects
                unique_subjects = {}
                for email_data in st.session_state.emails:
                    subject = email_data['subject']
                    if subject not in unique_subjects:
                        unique_subjects[subject] = []
                    unique_subjects[subject].append(email_data)

                total_request = len(unique_subjects)
                st.write(total_request)
                # Sort subjects by the most recent email date
                sorted_subjects = sorted(unique_subjects.keys(), key=lambda s: max(email['date'] for email in unique_subjects[s] if email['date']), reverse=True)
                
                # Display unique subjects
                selected_subject = st.selectbox("Select a unique request", sorted_subjects)
                
                # Display responses related to the selected subject
                related_emails = unique_subjects[selected_subject]
                
                # Sort related emails by date in descending order
                related_emails_sorted = sorted(related_emails, key=lambda e: e['date'] if e['date'] else datetime.min)
                
                email_options = [f"{email['subject']} - {email['date'].strftime('%Y-%m-%d %H:%M:%S') if email['date'] else 'No Date'}" for email in related_emails_sorted]
                selected_email_index = st.selectbox("Select a response", email_options)
                selected_email = related_emails_sorted[email_options.index(selected_email_index)]
    
                st.session_state.fetched_subject = st.text_area("Client Subject:", selected_email['subject'], height=50)
                st.session_state.fetched_content = st.text_area("Client Content:", selected_email['body'], height=500)
        
        else:
            st.session_state.fetched_subject = st.text_area("Client Subject:",st.session_state.fetched_subject,height = 50)
            st.session_state.fetched_content = st.text_area("Client Content:",st.session_state.fetched_content, height=500)
        col1, col3 = st.columns([1, 1])
        with col3:
            if st.button("Next"):
                if st.session_state.fetched_content and st.session_state.fetched_subject:
                    st.session_state.gmail_fetched = True
                    navigate_to_step(3)
                else:
                    st.error("Insert Client Request or Fetch Request from Gmail.")
        with col1:
            if st.button("Previous"):
                navigate_to_step(1)
    
    # Step 3: Evaluate and provide feedback
    elif st.session_state.step == 3:
        # if st.sidebar.button("Step 1: Upload SOP"):
        #     navigate_to_step(1)
        # if st.sidebar.button("Step 2: Client Request"):
        #     navigate_to_step(2)
        st.title("Step 3: Type your content to evaluate")
        if not st.session_state.user_gmail:
            st.session_state.user_gmail = st.text_input('User Mail',st.session_state.user_gmail)
        st.session_state.user_subject = st.text_area("Subject:",st.session_state.user_subject, height=50)
        st.session_state.user_msg= st.text_area("Message:",st.session_state.user_msg, height=400)
        st.session_state.user_input = f"subject: {st.session_state.user_subject} content:{st.session_state.user_msg}"
        
        if st.button("Evaluate"):
            if len(st.session_state.user_input) < 20:
                st.error("Insufficient Information")
            else:
                st.session_state.evaluation_count = st.session_state.evaluation_count + 1
                prompt = f"""
                As a Quality Analyst, your task is to meticulously evaluate a user's response to a client email based on our Standard Operating Procedure (SOP) for email communication. The client email outlines an issue or concern they are experiencing with our product. Your evaluation involves identifying the specific problem mentioned by the client and ensuring the response adheres to our SOP. Follow these steps:
                
                ### SOP Content
                {st.session_state.sop_content}
                
                ### Client Email
                **Subject:** {st.session_state.fetched_subject}
                **Content:** {st.session_state.fetched_content}
                
                ### Evaluation Task
                
                1. Client's Issue:
                   - Clearly identify the specific problem or concern mentioned by the client in their email.
                
                2. Constructive Feedback:
                   - Provide actionable feedback aimed at improving future responses.
                   - Ensure feedback is specific and provides clear examples where applicable.
                
                3. Criteria Instruction in SOP:
                   - Go through the SOP.
                   - Provide the instruction for each criterion from the SOP.
                
                4. Evaluation Based on SOP:
                   - For each criterion in the SOP, provide a mark (out of 10) with a reason for the score within 25 words.
                   - Present the criteria in a 2D list format:
                   
                     | Criteria | Mark (out of 10) | Reason |
                     | --- | --- | --- |
                     | [Criteria 1] | [Score] | [Reason] |
                     | [Criteria 2] | [Score] | [Reason] |
                     | ... | ... | ... |
                   - Give the mark as a single integer.
                
                5. Suggested Alternatives:
                   - Suggest better alternative email content, fully structured with subject and body, that aligns with the SOP and addresses the client's concern effectively.
                """
    
                try:
                    completion = client.chat.completions.create(
                        messages=[
                            {"role": "system", "content": prompt},
                            {"role": "user", "content": st.session_state.user_input}
                        ],
                        model="llama3-70b-8192",
                        temperature=0,
                    )
                    st.session_state.feedback = completion.choices[0].message.content
                except Exception as e:
                    st.error(f"An error occurred: {e}")
        if st.session_state.feedback:
            # st.text_area('chheck', st.session_state.feedback,height=600)
            feedback_parts = st.session_state.feedback.split("Suggested Alternatives:")
            feedback_text = feedback_parts[0].strip()
            # st.write(st.secrets)

            # conn = st.connection("gsheets", type=GSheetsConnection)
            if feedback_text:
                try:
                    feedback_criteria, sop_evaluation = process_feedback(feedback_text)
                    feedback, criteria = process_criteria(feedback_criteria)
                    left, right = st.columns([2, 1])
                    # with left:
                    st.subheader('Feedback and Criteria')
                    st.write(feedback)
    
                    st.subheader("Criteria Instruction In SOP:")
                    st.write(criteria)
                    
                    st.subheader('Evaluation Based on SOP')
                    scope = ["https://www.googleapis.com/auth/spreadsheets"]
                    
                    # Load credentials from Streamlit secrets
                    credentials = Credentials.from_service_account_info(
                        st.secrets["gcp_service_account"],
                        scopes=scope
                    )
                    # st.write("hello there")
                except Exception as e:
                    st.error(f"An error occurred: {e}")
                try:
                    # Create a client to interact with the Google Drive API
                    sop_client = gspread.authorize(credentials)
                    # st.write("hello tere")
                except Exception as e:
                    st.error(f"An error occurred: {e}")

                try:
                    # Open the Google Sheet
                    sheet = sop_client.open_by_key("1h9Mfg-VRZKm9YB9dhg5iguX_Yfg01BCyl2qzpvGPMcU").worksheet("Employee Performance 1")
                    # st.write('hellololo')
                    df = parse_sop_evaluation(sop_evaluation)
                    df = df.drop(0)
                    # st.dataframe(df)
                    # conn.write(df, sheet="Employee Performance", mode="append")
                        # Convert DataFrame to a list of lists for appending
                    # st.write(df)
                    # st.write('table')
                    data = sheet.get_all_values()
                    criteria = df['Criteria'].tolist()
                    marks = df['Mark (out of 10)'].tolist()
                    marks = []
                    for mark in df['Mark (out of 10)']:
                        try:
                            marks.append(int(mark))
                        except ValueError:
                            # Skip non-numeric values like 'N/A'
                            continue
                    criteria_marks_dict = dict(zip(criteria, marks))
                    # st.write(data[0])
                    reason = ' | '.join(df['Reason'].values)
                    # st.write(reason)
                    # reason = reason.replace("'", "\\'").replace('\n', ' ')
                    # reason = str(reason)
                    # st.write(len(data))
                    batch_data = []
                    if data and len(data) > 1:
                        columns = data[0]
                        batch_data = [dict(zip(columns, row)) for row in data[1:]]
                    else:
                        columns = ['User', 'Time', 'Client_Request_subject', 'client_Request', 'User_Email', 'Reason']

                    for c in criteria:
                        if c not in columns:
                            columns.insert(-1,c)
                    # Create a new DataFrame with criteria as columns and marks as a single row
                    transformed_data = {
                        'User': st.session_state.user_gmail,
                        'Time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        'Client_Request_subject': st.session_state.fetched_subject,
                        'client_Request' : st.session_state.fetched_content,
                        'User_Email':st.session_state.user_input,
                        'Reason': reason
                    }
                    # st.write(transformed_data)
                    modi_transformed_data = transformed_data | criteria_marks_dict
                    # st.write(modi_transformed_data)
                    batch_data.append(modi_transformed_data)
                    pre_transformed_df = pd.DataFrame(batch_data, columns = columns)
                    # st.write(pre_transformed_df)
                    transformed_df = pre_transformed_df.fillna('')
                    data_to_append = transformed_df.values.tolist()
                    # st.write(data)
                    # if len(data) > 1:
                    #     data_to_append = df.values.tolist()
                    # else:
                    #     data_to_append = [df.columns.tolist()] + df.values.tolist()
                    # st.write(data_to_append)
            
                    # if st.session_state.evaluation_count > st.session_state.gsheet_count:
                    #     sheet.append_rows(data_to_append, value_input_option="RAW")
                    #     st.session_state.gsheet_count = st.session_state.gsheet_count + 1
                    # st.success("Data appended to Google Sheets successfully!")
                    st.table(df)
                except Exception as e:
                    st.session_state.evaluation_count = st.session_state.evaluation_count - 1
                    st.error(f"An error occurred: {e}")
                
                # st.write(data)
                # if len(data) > 0:
                #     data_to_append = df.values.tolist()
                # else:
                # data_to_append = [df.columns.tolist()] + df.values.tolist()
                        
                if st.session_state.evaluation_count > st.session_state.gsheet_count:
                    sheet.clear()
                    try:
                        sheet.append_rows([columns] + data_to_append)
                    except Exception as e:
                        st.write(e)
                    st.session_state.gsheet_count = st.session_state.gsheet_count + 1
                    st.success("Record Updated")
            
            suggested_alternatives_text = feedback_parts[1].strip()
            subject_start = suggested_alternatives_text.find("Subject:**")
            subject_end = suggested_alternatives_text.find("\n\n", subject_start)
            st.session_state.subject = suggested_alternatives_text[subject_start + len("Subject:**"):subject_end].strip()
            content_start = suggested_alternatives_text.find("Dear")
            if content_start == -1:
                content_start = suggested_alternatives_text.find("Hi")
            # content_start = subject_end + 2
            st.session_state.content = suggested_alternatives_text[content_start:].strip()
    
            st.title("Suggested Alternatives")
            st.session_state.subject = st.text_area("Subject", st.session_state.subject, height=100)
            st.session_state.content = st.text_area("Content", st.session_state.content, height=300)
            # st.text_area("Content", st.session_state.feedback, height=300)
            # st.text_area("feedback",st.session_state.feedback,height = 500)
            
            # if st.button("Send Email",key = 'process_start'):
        # if st.button("Next"):
            
    
        # if st.button("Previous"):
        # st.text_area("feedback",st.session_state.feedback,height = 500)
        col1, col3 = st.columns([1, 1])
        with col1:
            if st.button("Previous"):
                navigate_to_step(2)
        
        with col3:
            if st.button("Next"):
                if st.session_state.feedback:
                    navigate_to_step(4)
                else:
                    st.error("Evaluate before sending the mail.")

    elif st.session_state.step == 4:
        # if st.sidebar.button("Step 1: Upload SOP"):
        #     navigate_to_step(1)
        # if st.sidebar.button("Step 2: Client Request"):
        #     navigate_to_step(2)
        # if st.sidebar.button("Step 3: Evaluation and feedback"):
        #     navigate_to_step(3)
        # try:
        send_reply_email()
        # except Exception as e:
        #     st.error("please Authenticate your mail")
        col1, col3 = st.columns([1,1])
        with col1:
            if st.button("Previous"):
                navigate_to_step(3)
    
    # Navigation buttons

    # if st.session_state.step != 1:
    #     if st.sidebar.button("Step 1: Upload SOP"):
    #         navigate_to_step(1)
    # if st.session_state.step != 2:
    #     if st.sidebar.button("Step 2: Client Request"):
    #         navigate_to_step(2)
    # if st.session_state.step != 3:
    #     if st.sidebar.button("Step 3: Evaluation and feedback"):
    #         navigate_to_step(3)
    # if st.session_state.step != 4:



def sop_creator(client):
    # st.checkbox("Grammer")
    st.session_state.department = st.text_input("Industry:",st.session_state.department)
    st.session_state.purpose = st.text_input("Purpose of the SOP:", st.session_state.purpose)
    st.session_state.procedure = st.text_area("Process Details",st.session_state.procedure, height = 400)
    st.session_state.criteria = st.text_area("Metrics / criteria to be included ",st.session_state.criteria, height = 200)
    if st.button("Create SOP"):
        if st.session_state.department and st.session_state.purpose and st.session_state.criteria:
            prompt =f"""
                You are an SOP (Standard Operating Procedure) engineer tasked with creating a comprehensive SOP for the {st.session_state.department}. 
                The purpose of this SOP is to {st.session_state.purpose} and the workflow given by the user. Ensure that the SOP adheres to the following criteria:
                
                Introduction:
                
                Provide a clear overview of the SOP.
                State the objectives and scope.
                Mention the target audience and their roles.
                Purpose:
                
                Explain the specific goals and intended outcomes of this SOP.
                Detail how it aligns with the department's overall mission and objectives.
                Definitions:
                
                Define any technical terms, acronyms, or jargon used within the SOP.
                Include a glossary if necessary.
                Procedures:
                
                Step-by-step instructions for completing the tasks outlined in the SOP.
                Include detailed sub-steps, decision points, and potential variations.
                Use diagrams, flowcharts, or checklists to enhance clarity.
                Roles and Responsibilities:
                
                Clearly define the roles and responsibilities of each individual or team involved.
                Specify any required qualifications or training.
                Materials and Equipment:
                
                List all necessary materials, tools, and equipment needed to perform the procedures.
                Include specifications or standards if applicable.
                Safety and Compliance:
                
                Highlight any safety precautions or regulatory compliance requirements.
                Include guidelines for emergency procedures and reporting incidents.
                Documentation and Records:
                
                Specify the type of documentation required at each step.
                Detail how records should be maintained, stored, and accessed.
                Quality Assurance:
                
                Outline the measures in place to ensure the quality and consistency of the procedures.
                Include methods for auditing and continuous improvement.
                Criteria:
                
                Ensure the SOP meets the specific criteria mentioned by the user: {st.session_state.criteria}.
                Must Provide detailed instructions or guidelines on how to adhere to these criteria.
                Appendices:
                
                Add any supplementary information such as templates, sample forms, or additional references."""
            try:
                completion = client.chat.completions.create(
                messages=[
                                {"role": "system", "content": prompt},
                                {"role": "user", "content": f"workflow Procedures{st.session_state.procedure}"}
                            ],
                            model="llama3-70b-8192",
                            temperature=0,
                        )
                st.session_state.created_sop = completion.choices[0].message.content
            except Exception as e:
                st.error(f"An error occurred: {e}")
            st.session_state.created_sop = st.text_area(f"SOP for {st.session_state.purpose}",st.session_state.created_sop, height=500)
            file_name = st.text_input("Enter the name of the file to save (e.g., improved_text.txt):", "sop.txt")
    
                # Button to save the edited content to a new file
            if st.session_state.created_sop:
                    # Save the edited content to a new file
                st.download_button(
                        label="Download SOP file",
                        data=st.session_state.created_sop.encode('utf-8'),
                        file_name=file_name,
                        mime="text/plain"
                    )
                st.success(f"File '{file_name}' processed and ready for download.")
        else:
            st.error("Fill all the neccessary field")
    
    

def main():
    # st.set_page_config(layout="wide")
    
    st.markdown("""
        <style>
        div.stButton > button:first-child {
            width: 100%;
        }
        </style>
        """, unsafe_allow_html=True)



    # st.session_state.col1, st.session_state.col2, st.session_state.col3 = st.columns([1, 4, 1])
    
    if 'sop_uploaded' not in st.session_state:
        st.session_state.sop_uploaded = False
    if 'sop_content' not in st.session_state:
        st.session_state.sop_content = ''
    if 'gmail_fetched' not in st.session_state:
        st.session_state.gmail_fetched = False
    if 'gmail_content' not in st.session_state:
        st.session_state.gmail_content = ""
    if 'user_input' not in st.session_state:
        st.session_state.user_input = ""
    if 'gmail_send' not in st.session_state:
        st.session_state.gmail_send = False
    if 'feedback' not in st.session_state:
        st.session_state.feedback = ""
    if 'gmail_sender' not in st.session_state:
        st.session_state.gmail_sender = ""
    if 'user_gmail' not in st.session_state:
        st.session_state.user_gmail = ""
    if 'auth_url' not in st.session_state:
        st.session_state.auth_url = ""
    if 'subject' not in st.session_state:
        st.session_state.subject = ""
    if 'content' not in st.session_state:
        st.session_state.content = ""
    if 'msg_id' not in st.session_state:
        st.session_state.msg_id = ""
    if 'password' not in st.session_state:
        st.session_state.password = ""
    if 'department' not in st.session_state:
        st.session_state.department = ""
    if 'criteria' not in st.session_state:
        st.session_state.criteria = ""
    if 'purpose' not in st.session_state:
        st.session_state.purpose = ""
    if 'fetched_sender_gmail' not in st.session_state:
        st.session_state.fetched_sender_gmail = ""
    if 'fetched_subject' not in st.session_state:
        st.session_state.fetched_subject = ""
    if 'fetched_content' not in st.session_state:
        st.session_state.fetched_content = ""
    if 'default_sop_content' not in st.session_state:
        st.session_state.default_sop_content = ""
    if 'user_subject' not in st.session_state:
        st.session_state.user_subject = ""
    if 'user_msg' not in st.session_state:
        st.session_state.user_msg = ""
    if 'step_1' not in st.session_state:
        st.session_state.step_1 = True
    if 'step_2' not in st.session_state:
        st.session_state.step_2 = False
    if 'step_3' not in st.session_state:
        st.session_state.step_3 = False
    if 'Authenticate_check' not in st.session_state:
        st.session_state.Authenticate_check = False
    if 'step' not in st.session_state:
        st.session_state.step = 1
    if 'evaluation_count' not in st.session_state:
        st.session_state.evaluation_count = 0
    if 'email_count_total' not in st.session_state:
        st.session_state.email_count_total = 0
    if 'gsheet_count' not in st.session_state:
        st.session_state.gsheet_count = 0
    if 'sop_created' not in st.session_state:
        st.session_state.created_sop = ""
    if 'procedure' not in st.session_state:
        st.session_state.procedure = ""

    # gsheet_id = "1WWGaGc-rVpMYhUjsxDuYDELynzQlq5XKVv3kDU1DuVU"
    # workbook = get_workbook(gsheet_id)
    # try:
    #     sheet = workbook.worksheet('hello')
    # except gspread.exceptions.WorksheetNotFound:
    #     sheet = workbook.add_worksheet(title='hello', rows="1000", cols="200")
    
    client = Groq(api_key=st.secrets["API_KEY"])
    option = st.selectbox("Choose the tool", ("Evaluator","SOP creator",), index=None, placeholder='Choose an option')
    if option == "Evaluator":
        evaluator(client)
    elif option == "SOP creator":
        sop_creator(client)

if __name__ == "__main__":
    main()
